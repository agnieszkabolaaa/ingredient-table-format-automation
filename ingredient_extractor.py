#!/usr/bin/env python3
"""
Automated system for extracting and formatting cosmetic ingredient data from Word documents.
Processes a messy base document with multiple ingredients and creates standardized
single-ingredient documents with consistent table structure.
"""

import os
import re
import unicodedata
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx2pdf import convert



class IngredientExtractor:
    """Extracts ingredient data from a Word document."""

    def __init__(self, base_doc_path: str, template_doc_path: str, output_dir: Optional[str] = None):
        """
        Initialize the extractor with base and template documents.

        Args:
            base_doc_path: Path to the messy base document
            template_doc_path: Path to the template document with standardized format
            output_dir: Optional output directory for generated documents
        """
        self.base_doc_path = self._resolve_existing_document(Path(base_doc_path))
        self.template_doc_path = self._resolve_existing_document(Path(template_doc_path))
        self.base_doc = Document(str(self.base_doc_path))
        self.template_doc = Document(str(self.template_doc_path))
        self.ingredients: List[Dict] = []
        self.output_dir = Path(output_dir or "output_ingredients")
        self.output_dir.mkdir(exist_ok=True, parents=True)

    def _resolve_existing_document(self, doc_path: Path) -> Path:
        """Resolve a document path even when the filename uses a different unicode normalization."""
        if doc_path.exists():
            return doc_path

        search_dirs = [Path.cwd(), Path(__file__).resolve().parent]
        if doc_path.parent and doc_path.parent.exists():
            search_dirs.insert(0, doc_path.parent)

        normalized_target = self._normalize_text(doc_path.name)
        for directory in search_dirs:
            if not directory.exists():
                continue
            for candidate in directory.iterdir():
                if not candidate.is_file() or candidate.suffix.lower() != ".docx":
                    continue
                if self._normalize_text(candidate.name) == normalized_target:
                    return candidate

        return doc_path

    @staticmethod
    def _normalize_text(value: str) -> str:
        return unicodedata.normalize("NFC", value).casefold()

    def extract_ingredient_sections(self) -> List[Dict]:
        """
        Extract ingredient sections from the base document.
        Identifies ingredient names and their associated tables.

        Returns:
            List of dictionaries containing ingredient data.
        """
        ingredients_by_name: Dict[str, Dict] = {}
        current_ingredients: List[Dict] = []

        for table in self.base_doc.tables:
            table_data = self._extract_table_data(table)
            ingredients = self._extract_ingredient_names(table_data)



            if ingredients:
                current_ingredients = ingredients

                for ing in current_ingredients:

                    ingredient_name = ing["name"]
                    column = ing["column"]
                    if ingredient_name not in ingredients_by_name:
                        ingredients_by_name[ingredient_name] = {
                            "name": ingredient_name,
                            "column": column,
                            "content": [],
                            "tables": [],
                            "paragraphs": 0,
                        }
                    ingredients_by_name[ingredient_name]["tables"].append(table_data)
            elif current_ingredients:

                for ing in current_ingredients:

                    ingredient_name = ing["name"]
                    if ingredient_name not in ingredients_by_name:
                        ingredients_by_name[ingredient_name] = {
                            "name": ingredient_name,
                            "column": ing["column"],
                            "content": [],
                            "tables": [],
                            "paragraphs": 0,
                        }
                    ingredients_by_name[ingredient_name]["tables"].append(table_data)
                  
        self.ingredients = list(ingredients_by_name.values())
        return self.ingredients

    def attach_tables_to_ingredients(self) -> None:
        """Ensure each ingredient has its tables attached after extraction."""
        if not self.ingredients:
            self.extract_ingredient_sections()

        tables_by_ingredient = self.extract_tables_by_ingredient()
        for ingredient in self.ingredients:
            ingredient["tables"] = tables_by_ingredient.get(ingredient["name"], [])
            ingredient["paragraphs"] = len(ingredient.get("content", []))

    def extract_tables_by_ingredient(self) -> Dict[str, List]:
        """Extract tables from the base document organized by ingredient."""
        self.extract_ingredient_sections()
        return {ingredient["name"]: ingredient.get("tables", []) for ingredient in self.ingredients}

    def _extract_table_data(self, table) -> List[List[str]]:
        """Convert a table to a list of lists containing cell data."""
        return [[cell.text.strip() for cell in row.cells] for row in table.rows]

    def _extract_ingredient_names(self, table_data: List[List[str]]) -> List[Dict[str, Any]]:
        for row in table_data:
            if not row:
                continue

            if "inci name" not in row[0].strip().lower():
                continue

            ingredients = []

            for col_idx, cell in enumerate(row[1:], start=1):
                cleaned = self._clean_ingredient_name(cell)
                if cleaned:
                    ingredients.append({
                        "name": cleaned,
                        "column": col_idx
                    })

            return ingredients
        return []

    def _clean_ingredient_name(self, value: str) -> Optional[str]:
        """Normalize a candidate ingredient name from a table cell."""
        candidate = re.sub(r"\s+", " ", value).strip().strip(":")
        if not candidate:
            return None
        lowered = candidate.lower()
        if lowered in {"inci name", "chemical (iupac) and common name", "toxicity profile", "toxicological profile", "toxicology profile"}:
            return None
        return candidate

    def create_standardized_document(self, ingredient_name: str, ingredient_data: Dict) -> Document:
        """Create a standardized document that preserves the template structure and formatting."""
        new_doc = Document(str(self.template_doc_path))
        self._populate_document(new_doc, ingredient_name, ingredient_data)
        return new_doc

    def _populate_document(self, doc: Document, ingredient_name: str, ingredient_data: Dict) -> None:
        """Populate the template table in place while preserving its existing formatting."""
        values_by_label = self._build_template_row_values(
            ingredient_name,
            ingredient_data
        )
        table = self._find_template_table(doc)
        if table is None:
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Grid Accent 1"

        for row in table.rows:
            label_cell = row.cells[0].text if row.cells else ""
            label = self._map_template_label(label_cell)
            if not label:
                continue

            value = values_by_label.get(label, "")
            self._fill_template_row(row, label, value)

            if label == "Chemical (IUPAC) and common name" and len(row.cells) > 1:
                row.cells[1].text = value

        if table.rows:
            first_row = table.rows[0]
            if first_row.cells and len(first_row.cells) > 1:
                first_row.cells[1].text = ingredient_name

    def _find_template_table(self, doc: Document):
        """Return the first table in a document so we can populate the template structure."""
        for table in doc.tables:
            if table.rows and table.columns:
                return table
        return None

    def _fill_template_row(self, row, label: str, value: str) -> None:
        """Populate the value cells in a template row without changing the label cells."""
        if len(row.cells) <= 1:
            return

        # Wiersze ze scalonymi komórkami
        if label in {
            "Mutagenicity/genotoxicity",
            "Carcinogenicity",
            "Reproductive toxicity",
        }:
            label_tc = row.cells[0]._tc

            for cell in row.cells[1:]:
                if cell._tc != label_tc:
                    cell.text = value
                    break
            return

        if label == "INCI name":
            start_index = 1
        elif len(row.cells) >= 4:
            start_index = 2
        else:
            start_index = 1

        for index in range(start_index, len(row.cells)):
            row.cells[index].text = value

    def _build_template_row_values(
        self,
        ingredient_name: str,
        ingredient_data: Dict[str, Any]
    ) -> Dict[str, str]:
        """Map ingredient data from the source tables to the template row labels."""
        row_values: Dict[str, str] = {"INCI name": ingredient_name}
        tables = ingredient_data["tables"]
        column = ingredient_data["column"]

        for table_data in tables:
            for row_data in table_data:
                if len(row_data) < 2:
                    continue

                raw_label = row_data[0].strip()
                if not raw_label:
                    continue

                target_label = self._map_template_label(raw_label)
                if not target_label:
                    continue

                value = self._extract_table_value(row_data, column)
                if value:
                    row_values[target_label] = value

        return row_values

    def _map_template_label(self, raw_label: str) -> Optional[str]:
        """Translate source-table labels into the template row labels."""

        normalized = self._normalize_label(raw_label)
        if normalized in {"inci name"}:
            return "INCI name"
        if normalized in {"chemical iupac and common name", "chemical iupac and common name:"}:
            return "Chemical (IUPAC) and common name"
        if normalized in {"toxicity profile", "toxicological profile", "toxicology profile"}:
            return "Toxicity profile"
        if normalized in {"acute toxicity oral"}:
            return "Acute toxicity (oral)"
        if normalized in {"acute toxicity dermal"}:
            return "Acute toxicity (dermal)"
        if normalized in {"acute toxicity inhalation"}:
            return "Acute toxicity (inhalation)"
        if normalized in {"skin irritation"}:
            return "Skin irritation"
        if normalized in {"eye irritation", "mucus membrane irritation"}:
            return "Eye irritation"
        if normalized in {"skin sensitisation", "skin sensitization", "sensitization", "sensitisation"}:
            return "Skin sensitisation"
        if normalized in {"dermal absorption"}:
            return "Dermal absorption"
        if normalized.startswith("repeated dose toxicity"):
            return "Repeated dose toxicity"
        if normalized in {"mutagenicity genotoxicity"}:
            return "Mutagenicity/genotoxicity"
        if normalized in {"carcinogenicity"}:
            return "Carcinogenicity"
        if normalized in {"reproductive toxicity"}:
            return "Reproductive toxicity"
        if normalized in {"toxicokinetics"}:
            return "Toxicokinetics"
        if normalized in {"phototoxicity"}:
            return "Phototoxicity"
        if normalized in {"human data"}:
            return "Human data"
        if normalized.startswith("source"):
            return "Source"
        

        return None

    def _normalize_label(self, label: str) -> str:
        """Normalize labels so source rows can be matched to the template rows."""
        return re.sub(r"[^a-z0-9]+", " ", label.lower()).strip()

    def _extract_table_value(self, row_data: List[str], column: int) -> str:

         # how many non-empty values are there in the row after the label
        values = [c.strip() for c in row_data[1:] if c.strip()]

        if not values:
            return ""

        # one value -> common for all
        if len(values) == 1:
            return values[0]

        # multiple values -> use the one in the same column as the ingredient name
        if column < len(row_data):
            return row_data[column].strip()

        return values[0]

    def _add_formatted_table(self, doc: Document, table_data: List[List[str]]) -> None:
        """Add a formatted table to the document."""
        if not table_data:
            return

        rows = len(table_data)
        cols = max(len(row) for row in table_data) if table_data else 0
        if rows == 0 or cols == 0:
            return

        table = doc.add_table(rows=rows, cols=cols)
        table.style = "Light Grid Accent 1"

        for i, row_data in enumerate(table_data):
            for j, cell_value in enumerate(row_data):
                if j >= len(table.rows[i].cells):
                    continue
                table.rows[i].cells[j].text = str(cell_value)

    def process_all_ingredients(self) -> None:
        """Process all ingredients and create separate documents."""
        print(f"Processing {len(self.ingredients)} ingredients...")

        for ingredient in self.ingredients:
            ingredient_name = ingredient["name"]
            print(f"  Processing: {ingredient_name}")

            ingredient_data = {
                "name": ingredient_name,
                "column": ingredient["column"],
                "content": ingredient.get("content", []),
                "tables": ingredient.get("tables", []),
            }

            new_doc = self.create_standardized_document(ingredient_name, ingredient_data)
            output_filename = self._sanitize_filename(ingredient_name)
            output_path = self.output_dir / f"{output_filename}.docx"

            try:
                new_doc.save(str(output_path))
                pdf_path = output_path.with_suffix(".pdf")
                convert(str(output_path), str(pdf_path))
                print(f"✓ Saved:")
                print(f"   DOCX: {output_path}")
                print(f"   PDF : {pdf_path}")
            except Exception as exc:
                print(f"    ✗ Error saving {output_filename}: {str(exc)}")

    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize ingredient name for use as filename."""
        
        invalid_chars = r'[<>:"/\\|?*]'
        sanitized = re.sub(invalid_chars, "_", filename)
        sanitized = sanitized.replace(" ", "_").upper()
        return sanitized[:255]

    def print_summary(self) -> None:
        """Print a summary of extracted ingredients."""
        print("\n" + "=" * 60)
        print("EXTRACTION SUMMARY")
        print("=" * 60)
        print(f"Total ingredients found: {len(self.ingredients)}\n")

        for i, ingredient in enumerate(self.ingredients, 1):
            print(f"{i}. {ingredient['name']}")
            print(f"   - Content paragraphs: {ingredient['paragraphs']}")
            print(f"   - Tables: {len(ingredient.get('tables', []))}")

        print("=" * 60 + "\n")


def main() -> None:
    """Main execution function."""

    import sys

    template_doc = "00_WZÓR.docx"

    if getattr(sys, 'frozen', False):
        repo_root = Path(sys.executable).resolve().parent
    else:
        repo_root = Path(__file__).resolve().parent

    os.chdir(repo_root)
    template_path = repo_root / template_doc

    if not template_path.exists():
        normalized_template = IngredientExtractor._normalize_text(template_doc)
        for candidate in repo_root.glob("*.docx"):
            if IngredientExtractor._normalize_text(candidate.name) == normalized_template:
                template_path = candidate
                break

    if not template_path.exists():
        print(f"Error: {template_doc} not found")
        return

    # Znajdź wszystkie dokumenty źródłowe
    base_documents = []

    for candidate in repo_root.glob("*.docx"):

        normalized = IngredientExtractor._normalize_text(candidate.name)

        # pomiń szablon
        if normalized == IngredientExtractor._normalize_text(template_doc):
            continue

        # pomiń pliki tymczasowe Worda
        if candidate.name.startswith("~$"):
            continue

        # pomiń wygenerowane dokumenty
        if candidate.parent.name == "output_ingredients":
            continue

        base_documents.append(candidate)

    if not base_documents:
        print("No source .docx files found.")
        return

    print(f"Found {len(base_documents)} source document(s):")
    for doc in base_documents:
        print(f" - {doc.name}")

    print()

    # Przetwarzanie każdego pliku
    for base_doc in base_documents:

        print("=" * 70)
        print(f"Processing file: {base_doc.name}")
        print("=" * 70)

        extractor = IngredientExtractor(
            str(base_doc),
            str(template_path)
        )

        print("Extracting ingredients...")
        extractor.extract_ingredient_sections()

        print("Extracting tables...")
        extractor.attach_tables_to_ingredients()

        extractor.print_summary()

        print("Creating standardized documents...")
        extractor.process_all_ingredients()

    print(f"\n✓ All documents saved to: {repo_root / 'output_ingredients'}")


if __name__ == "__main__":
    main()